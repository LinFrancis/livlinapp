"""Generador Word — LivLin · ERP/HRP · Todas las secciones."""
import io, base64, json
from datetime import datetime
from pathlib import Path
from utils.petal_content import (
    PETAL_DESC, PETAL_ICONS, IPR_SCALE, IPR_WHAT_IS, IPR_OBS_VS_POT,
    LIVLIN_URL, LIVLIN_TAGLINE, LIVLIN_DESC, LIVLIN_MODULES,
    LIVLIN_SERVICES_PITCH, LIVLIN_CLOSING, REGENERATION_FRAMEWORK, GLOBAL_REFS,
    TAO_REFLEXION_SHORT, TAO_INVITACION
)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

C1 = RGBColor(0x1B,0x43,0x32) if DOCX_OK else None
C2 = RGBColor(0x2D,0x6A,0x4F) if DOCX_OK else None
C3 = RGBColor(0x52,0xB7,0x88) if DOCX_OK else None
CA = RGBColor(0xE6,0x51,0x00) if DOCX_OK else None
CG = RGBColor(0x44,0x44,0x44) if DOCX_OK else None

def _shd(cell, hx):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),hx); pr.append(s)
def _H(d,t,lv=1,c=None,sz=None):
    p=d.add_heading(t,level=lv); r=p.runs[0] if p.runs else p.add_run(t)
    if c: r.font.color.rgb=c
    if sz: r.font.size=Pt(sz)
def _p(d,t,bold=False,c=None,sz=10,it=False,af=4):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(af)
    r=p.add_run(t); r.bold=bold; r.italic=it; r.font.size=Pt(sz)
    if c: r.font.color.rgb=c
def _box(d,t,sz=9,c=None):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t); r.font.size=Pt(sz); r.italic=True; r.font.color.rgb=c or C2
def _kv(d,rows):
    if not rows: return
    t=d.add_table(rows=len(rows),cols=2); t.style="Table Grid"
    for i,(l,v) in enumerate(rows):
        _shd(t.rows[i].cells[0],"D8F3DC"); _shd(t.rows[i].cells[1],"FFFFFF")
        for c in t.rows[i].cells: c.paragraphs[0].paragraph_format.space_after=Pt(0)
        r0=t.rows[i].cells[0].paragraphs[0].add_run(str(l)); r0.bold=True; r0.font.size=Pt(9); r0.font.color.rgb=C1
        r1=t.rows[i].cells[1].paragraphs[0].add_run(str(v)[:500] if v else "—"); r1.font.size=Pt(9); r1.font.color.rgb=C2
def _logo(d,w=1.5):
    try:
        from utils.logo_b64 import LOGO_B64; import os
        tp="/tmp/livlin_logo.png"
        if not os.path.exists(tp):
            with open(tp,"wb") as f: f.write(base64.b64decode(LOGO_B64))
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(tp,width=Inches(w))
    except: p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("🌿 LivLin"); r.bold=True; r.font.size=Pt(18)
def _sf(v):
    try: return float(v) if v else 0
    except: return 0

def generate_docx(data:dict)->bytes:
    if not DOCX_OK: raise ImportError("python-docx no instalado.")
    d=Document()
    for s in d.sections: s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

    from utils.scoring import (compute_erp, compute_hrp, compute_domain_scores, compute_domain_scores_total,
        compute_cross_module_score, CROSS_MODULE_DETAIL, score_label, brecha_valor, brecha_texto,
        _score_to_level, get_petal_interp, compute_synthesis_potentials, compute_synthesis_potentials_obs,
        get_interp_text, PETAL_ORDER, DIM_WHAT_MEASURES)
    erp=compute_erp(data); hrp=compute_hrp(data); brecha=brecha_valor(erp,hrp); brecha_txt=brecha_texto(erp,hrp)
    le,_=score_label(erp); lh,_=score_label(hrp)
    do=compute_domain_scores(data); dt=compute_domain_scores_total(data)
    cross=compute_cross_module_score(data); pe=compute_synthesis_potentials_obs(data); ph_=compute_synthesis_potentials(data)
    nombre=data.get("proyecto_nombre","Diagnóstico"); cliente=data.get("proyecto_cliente","—")
    petalos=[]
    try:
        jf=Path(__file__).parent.parent/"data"/"petalos_regeneracion_urbana.json"
        with open(jf,encoding="utf-8") as f: petalos=json.load(f)["petalos"]
    except: pass

    # ═══ PORTADA ═══
    d.add_paragraph(); _logo(d); d.add_paragraph()
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("INDAGACIÓN REGENERATIVA"); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=C1
    p2=d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run("Diagnóstico de Permacultura Regenerativa"); r2.font.size=Pt(12); r2.font.color.rgb=C2
    p3=d.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=p3.add_run(f"LivLin · {LIVLIN_TAGLINE} · www.livlin.cl"); r3.font.size=Pt(10); r3.font.color.rgb=C3; r3.italic=True
    d.add_paragraph()
    rows=[("🏡 Espacio",nombre),("👤 Contacto",cliente),("📍 Ciudad",data.get("proyecto_ciudad","")),
          ("📅 Fecha",data.get("proyecto_fecha","") or str(datetime.now())[:10])]
    for l,k in [("📏 Superficie (m²)","proyecto_superficie"),("👥 Habitantes","proyecto_habitantes"),
                ("🐾 Mascotas","proyecto_mascotas"),("🏠 Tipo","proyecto_tipo_espacio")]:
        v=data.get(k); 
        if v: rows.append((l,str(v)))
    _kv(d,[r for r in rows if r[1]])
    d.add_paragraph()
    _H(d,"ERP y HRP",2,C1,12)
    _kv(d,[("🌍 ERP",f"{erp}/10 — {le}"),("🌱 HRP",f"{hrp}/10 — {lh}"),("🌀 Brecha",f"{brecha} pts — {brecha_txt}")])
    d.add_page_break()

    # === TAO ===
    _H(d,"Tao de la Regeneracion",1,C1,14)
    _box(d,TAO_REFLEXION_SHORT,c=C2)
    try:
        from modules.tao import TAO_DIMENSIONES, get_tao_scores, get_tao_total, get_tao_label
        scores = get_tao_scores(data)
        total = get_tao_total(data)
        label = get_tao_label(total)
        _kv(d,[("Puntaje Total",f"{total}/25 -- {label}")])
        tr = []
        for dim in TAO_DIMENSIONES:
            v = scores.get(dim["id"], 0)
            if v > 0:
                opt_text = dim["opciones"][v-1] if 1 <= v <= 5 else ""
                tr.append((f"D{dim['icono']}: {dim['titulo']}", f"Nivel {v}/5 -- {opt_text}"))
        if tr: _kv(d,tr)
    except Exception:
        pass
    notas = data.get("tao_notas","")
    if notas: _kv(d,[("Notas Tao",notas)])
    # Cuenca
    cuenca = data.get("cuenca_nombre","")
    if cuenca:
        _H(d,"Cuenca Hidrografica",2,C2,11)
        tr = [("Cuenca",cuenca)]
        sc = data.get("subcuenca_nombre","")
        if sc: tr.append(("Subcuenca",sc))
        ssc = data.get("subsubcuenca_nombre","")
        if ssc: tr.append(("Subsubcuenca",ssc))
        _kv(d,tr)
    d.add_page_break()

    # ═══ ECOLOGÍA ═══
    _H(d,"Observación Ecológica del Sitio",1,C1,14)
    for sec,fields in [("Suelo",[("suelo_tipo","Tipo"),("suelo_compactacion","Compactación"),("suelo_materia_organica","Mat. orgánica"),
        ("suelo_drenaje","Drenaje"),("suelo_color","Color"),("suelo_olor","Olor"),("suelo_notas","Notas")]),
        ("Flujos",[("sol_horas","Horas sol/día"),("sol_horas_invierno","Sol invierno"),("sol_horas_verano","Sol verano"),
        ("sol_orientacion","Orientación"),("sol_zonas_max","Zonas soleadas"),("sol_sombra_perm","Sombra permanente"),
        ("viento_direccion","Dirección viento"),("agua_flujo_lluvia","Flujo lluvia")]),
        ("Vegetación/Fauna",[("veg_tipos","Tipos vegetación"),("veg_especies","Especies"),("veg_invasoras","Invasoras"),
        ("fauna_lombrices","Lombrices"),("fauna_plagas","Plagas"),("fauna_aves_especies","Aves"),("eco_notas","Notas eco")]),
        ("Cultivo",[("cultivo_m2","m² cultivables"),("cultivo_m2_futuro","m² futuros"),("cultivo_produce_hoy","Produce hoy"),
        ("cultivo_frutales","Frutales"),("cultivo_plantas_actuales","Plantas actuales")])]:
        rows=[(l,str(data.get(k,""))) for k,l in fields if data.get(k)]
        if rows: _H(d,sec,2,C2,11); _kv(d,rows)
    # Climate
    lat=_sf(data.get("geo_lat")); lon=_sf(data.get("geo_lon")); elev=_sf(data.get("geo_elevation"))
    if lat and lon:
        _H(d,"Datos climáticos (Open-Meteo API)",2,C2,11)
        cr=[("Latitud",f"{lat:.4f}"),("Longitud",f"{lon:.4f}")]
        if elev>0: cr.append(("Elevación",f"{elev:.0f} m.s.n.m."))
        for l,k in [("Mes cálido","clima_mes_caluroso"),("Mes frío","clima_mes_frio"),("T° máx","clima_t_max_abs"),
                    ("T° mín","clima_t_min_abs"),("Precip. anual (mm)","agua_prec_anual")]:
            v=data.get(k);
            if v: cr.append((l,str(v)))
        _kv(d,cr)
    d.add_page_break()

    # ═══ SISTEMAS ═══
    _H(d,"Contexto, Agua, Energía y Materiales",1,C1,14)
    for sec,fields in [("🏘️ Contexto",[("ctx_ind_verde","Entorno verde"),("ctx_cuenca","Cuenca"),("ctx_vecinos","Vecinos"),
        ("ctx_participacion","Participación"),("ctx_distancia_parques","Dist. parque (m)")]),
        ("💧 Agua",[("agua_ind_general","Percepción"),("agua_fuente","Fuente"),("agua_captacion_lluvia","Captación"),
        ("agua_grises","Grises"),("agua_riego_sistema","Riego"),("agua_sequias","Sequías"),("agua_sequias_impacto","Impacto sequía")]),
        ("⚡ Energía",[("ene_ind_general","Percepción"),("ene_fuente","Fuente"),("ene_led","LED"),("ene_solar_interes","Solar"),
        ("ene_regleta","Regletas"),("ene_circuitos","Circuitos"),("ene_monitoreo","Monitoreo"),("ene_apagar_luces","Apagar luces"),
        ("ene_eficiencia_electrodom","Electrodom. eficientes")]),
        ("♻️ Materiales",[("res_ind_general","Percepción"),("res_compostan","Compostaje"),("res_compost_tipo","Tipo compost"),
        ("res_separan","Separación"),("res_reutilizan","Reutilización"),("res_intentos_fallidos","Intentos fallidos")])]:
        rows=[(l,str(data.get(k,""))) for k,l in fields if data.get(k) and str(data.get(k)) not in ["No registrado",""]]
        if rows: _H(d,sec,2,C2,11); _kv(d,rows)
    # Extra computed fields
    extra=[]
    parques=data.get("ctx_parques_lista",[])
    if isinstance(parques,list) and parques: extra.append(("Parques","; ".join(f"{p.get('nombre','')} ({p.get('dist',0)}m)" for p in parques)))
    actores=data.get("ctx_actores",[])
    if isinstance(actores,list) and actores: extra.append(("Actores","; ".join(f"{a.get('nombre','')} ({a.get('tipo','')})" for a in actores)))
    cl=_sf(data.get("agua_consumo_estimado_ldia"))
    if cl>0: extra.append(("Consumo agua",f"{cl:.0f} L/día"))
    lc=_sf(data.get("agua_litros_captacion_anual"))
    if lc>0: extra.append(("Captación lluvia",f"{lc:,.0f} L/año"))
    ck=_sf(data.get("ene_consumo_cuenta_kwh"))
    if ck>0: extra.append(("Cuenta eléctrica",f"{ck:.0f} kWh/mes"))
    kwh=_sf(data.get("ene_kwh_dia_calc"))
    if kwh>0: extra.append(("Consumo equipos",f"{kwh:.1f} kWh/día"))
    equipos=data.get("equipos_electricos",[])
    if isinstance(equipos,list) and equipos:
        extra.append(("Equipos","; ".join(f"{e.get('nombre','')} ({e.get('kwh_dia',0)} kWh/día)" for e in equipos)))
    sh=_sf(data.get("sol_horas"))
    if sh>0 and kwh>0: extra.append(("☀️ Solar",f"~{max(1,round(kwh/(0.1*sh)))} paneles 100W"))
    if extra: _kv(d,extra)
    d.add_page_break()

    # ═══ FLOR ═══
    _H(d,"Flor de la Permacultura — ERP / HRP",1,C1,14)
    _box(d,IPR_WHAT_IS,c=C2)
    for i,pn in enumerate(PETAL_ORDER):
        ic=PETAL_ICONS[i] if i<len(PETAL_ICONS) else "🌱"
        e=do.get(pn,0); h=dt.get(pn,0); le2,_=_score_to_level(e); lh2,_=_score_to_level(h)
        _H(d,f"{ic} {pn}",2,C2,11)
        _p(d,f"ERP: {e:.0f}/10 ({le2})  |  HRP: {h:.0f}/10 ({lh2})  |  Brecha: +{round(h-e,1)}",bold=True,c=C1,sz=9)
        pi=PETAL_DESC.get(pn,{})
        if pi.get("resumen"): _box(d,pi["resumen"],sz=8,c=CG)
        ie=get_petal_interp(pn,e,"erp"); ih=get_petal_interp(pn,h,"hrp")
        if ie: _box(d,f"Estado actual: {ie}",sz=9,c=C1)
        if ih: _box(d,f"Horizonte: {ih}",sz=9,c=CA)
        obs=data.get(f"petalo_{i}_obs",{}); new=data.get(f"petalo_{i}_pot_new",{})
        oo=data.get(f"petalo_{i}_otros_obs",[]); on=data.get(f"petalo_{i}_otros_new",[])
        or_=[(k.replace("_"," ").title()," · ".join(v)) for k,v in obs.items() if v]
        if oo: or_.append(("Otras"," · ".join(oo)))
        nr=[(k.replace("_"," ").title()," · ".join(v)) for k,v in new.items() if v]
        if on: nr.append(("Otras"," · ".join(on)))
        if or_: _p(d,"Prácticas observadas (ERP):",bold=True,c=C1,sz=9,af=2); _kv(d,or_)
        if nr: _p(d,"Prácticas potenciales (→ HRP):",bold=True,c=CA,sz=9,af=2); _kv(d,nr)
        d.add_paragraph()
    if cross:
        _H(d,"Sub-indicadores M2-6 (20% del ERP)",2,C2,11)
        _kv(d,[(f"{info.get('icono','')} {n}",f"{info['score']}/10 — {info['fuente']}") for n,info in cross.items()])
        d.add_paragraph()
        _p(d,"Variables y escalas:",bold=True,c=C1,sz=10)
        for n,det in CROSS_MODULE_DETAIL.items():
            _p(d,f"{det.get('icono','')} {n}: {det['formula']}",bold=True,c=C2,sz=9)
            for vn,vs in det["variables"]: _p(d,f"  • {vn}: {vs}",c=CG,sz=8,af=1)
    d.add_page_break()

    # ═══ 10 DIMENSIONES ═══
    _H(d,"10 Dimensiones de Análisis",1,C1,14)
    for dim in pe:
        e=pe[dim]; h_=ph_.get(dim,0); di=DIM_WHAT_MEASURES.get(dim,{})
        if isinstance(di,str): di={"que_mide":di,"icono":"📊","fuentes":""}
        le3,_=_score_to_level(e); lh3,_=_score_to_level(h_)
        _H(d,f"{di.get('icono','📊')} {dim}",2,C2,11)
        _p(d,f"ERP: {e}/10 ({le3})  |  HRP: {h_}/10 ({lh3})  |  Brecha: +{round(h_-e,1)}",bold=True,c=C1,sz=9)
        if di.get("que_mide"): _box(d,f"¿Qué mide? {di['que_mide']}",sz=8,c=CG)
        ie=get_interp_text(dim,e,"erp"); ih=get_interp_text(dim,h_,"hrp")
        if ie: _box(d,f"Estado actual: {ie}",sz=9,c=C1)
        if ih: _box(d,f"Horizonte: {ih}",sz=9,c=CA)
    d.add_page_break()

    # ═══ METODOLOGÍA ═══
    _H(d,"Metodología",1,C1,14)
    _p(d,"ERP = 80% MFP observado + 20% Sub-indicadores M2-6",bold=True,c=C1,sz=10)
    _p(d,"Puntuación MFP: 0=0 | 1-2=2 | 3-5=4 | 6-9=6 | 10-14=8 | 15+=10",c=CG,sz=9)
    _p(d,"HRP = 100% MFP proyectado (observado + potencial)",bold=True,c=C1,sz=10)
    _p(d,"Brecha = HRP − ERP → campo de acción",bold=True,c=CA,sz=10)
    _kv(d,[("0-2","Sin inicio"),("2-4","Semilla"),("4-6","Brote"),("6-8","Crecimiento"),("8-10","Abundancia")])
    d.add_page_break()

    # ═══ SÍNTESIS ═══
    _H(d,"Síntesis y Plan de Acción",1,C1,14)
    for k,l in [("sint_fortalezas","💚 Fortalezas"),("sint_oportunidades","🌱 Oportunidades"),
                ("sint_limitaciones","⚡ Desafíos"),("sint_quick_wins","🎯 Primeros pasos")]:
        v=data.get(k,"");
        if v: _p(d,f"{l}: {v}",c=CG,sz=10)
    d.add_paragraph()
    for pk,fase in [("plan_inmediatas","⚡ Fase 1 (0–3 meses)"),("plan_estacionales","🌿 Fase 2 (3–12 meses)"),
                    ("plan_estructurales","🌳 Fase 3 (1–5 años)")]:
        v=data.get(pk)
        if v:
            _p(d,fase,bold=True,c=C2,sz=10)
            if isinstance(v,list):
                for it in v:
                    txt=it.get("titulo","") if isinstance(it,dict) else str(it)
                    if txt: p=d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); p.add_run(txt).font.size=Pt(9)
            else: _p(d,str(v),it=True,c=CG,sz=9)
    d.add_page_break()

    # ═══ CIERRE ═══
    _H(d,"LivLin — Tu aliado en la regeneración",1,C1,14)
    _logo(d); _box(d,LIVLIN_SERVICES_PITCH,sz=11,c=C1)
    d.add_paragraph()
    pw=d.add_paragraph(); pw.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rw=pw.add_run(f"🌿 {LIVLIN_TAGLINE} · www.livlin.cl"); rw.font.size=Pt(11); rw.font.color.rgb=C1; rw.bold=True
    d.add_page_break()

    # ═══ REFERENCIAS ═══
    _H(d,"Referencias",1,C1,14)
    for a,t,u in GLOBAL_REFS:
        _p(d,f"• {a} — {t}",c=C1,sz=9,af=2); _p(d,f"  🔗 {u}",it=True,c=C3,sz=8,af=4)
    pp=d.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=pp.add_run(f"LivLin · {str(datetime.now())[:10]}"); rr.font.size=Pt(8); rr.font.color.rgb=CG; rr.italic=True

    buf=io.BytesIO(); d.save(buf); buf.seek(0); return buf.read()
